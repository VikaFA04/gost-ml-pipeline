from __future__ import annotations

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.evaluation.format_regression_audit import (
    audit_negative_pair,
    audit_negative_directory,
    audits_to_frame,
    build_regression_predictions,
    best_positive_match,
    text_jaccard,
)


def write_docx(path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def test_text_jaccard_scores_overlap() -> None:
    assert text_jaccard({"a", "b"}, {"b", "c"}) == 1 / 3
    assert text_jaccard(set(), set()) == 1.0


def test_best_positive_match_uses_docx_text(tmp_path) -> None:
    positive_one = tmp_path / "1.docx"
    positive_two = tmp_path / "2.docx"
    negative = tmp_path / "negative.docx"
    write_docx(positive_one, ["alpha beta"])
    write_docx(positive_two, ["database sql query"])
    write_docx(negative, ["sql query database"])

    match, similarity = best_positive_match(negative, [positive_one, positive_two])

    assert match == positive_two
    assert similarity == 1.0


def test_audit_negative_pair_reports_before_after_diff(tmp_path) -> None:
    positive = tmp_path / "positive.docx"
    positive_doc = Document()
    positive_doc.add_paragraph("Paragraph")
    positive_doc.save(positive)

    negative = tmp_path / "negative.docx"
    negative_doc = Document()
    paragraph = negative_doc.add_paragraph("Paragraph")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    negative_doc.save(negative)

    audit = audit_negative_pair(
        positive_path=positive,
        negative_path=negative,
        workspace_dir=tmp_path / "workspace",
    )
    frame = audits_to_frame([audit])

    assert audit.before.changed_paragraphs == 1
    assert audit.after.changed_paragraphs == 1
    assert audit.formatter_summary["error"] == 0
    assert frame.loc[0, "negative"] == "negative.docx"
    assert frame.loc[0, "before_field_mismatches"] == 1
    assert frame.loc[0, "after_field_mismatches"] == 1
    assert frame.loc[0, "field_mismatch_delta"] == 0


def test_build_regression_predictions_preserves_bibliography_context(tmp_path) -> None:
    input_docx = tmp_path / "bibliography.docx"
    document = Document()
    document.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    document.add_paragraph("1 Теоретическая часть")
    document.add_paragraph("\tИсточник 1")
    document.add_paragraph("2 Практическая часть")
    document.add_paragraph("\tИсточник 2")
    document.save(input_docx)

    predictions_csv = tmp_path / "predictions.csv"
    build_regression_predictions(input_docx, predictions_csv)

    df = pd.read_csv(predictions_csv)
    assert df["postprocessed_label"].tolist() == [
        "bibliography_title",
        "title_section",
        "bibliography_item",
        "title_section",
        "bibliography_item",
    ]


def test_audit_negative_directory_skips_word_lock_files(tmp_path) -> None:
    positive_dir = tmp_path / "positive"
    negative_dir = tmp_path / "negative"
    workspace_dir = tmp_path / "workspace"
    positive_dir.mkdir()
    negative_dir.mkdir()

    write_docx(positive_dir / "positive.docx", ["Paragraph"])
    write_docx(negative_dir / "negative.docx", ["Paragraph"])
    (negative_dir / "~$negative.docx").write_text("not a real docx", encoding="utf-8")

    audits = audit_negative_directory(positive_dir, negative_dir, workspace_dir)

    assert [audit.negative_path.name for audit in audits] == ["negative.docx"]
