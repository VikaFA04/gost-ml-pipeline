"""Unit tests for src/rules/style_signatures.classify_style.

RED-state tests written in Wave 0 (phase 01-engine-guardrails-cohesion-audit).
Plan 02 (Wave 1) implements style_signatures.py so these turn GREEN.
"""

from __future__ import annotations

from types import SimpleNamespace

from docx import Document
from docx.shared import Pt, Cm

from src.rules.style_signatures import classify_style


def _paragraph_with_style_name(name: str):
    """Shim — classify_style only reads .style.name, so a SimpleNamespace suffices.

    Avoids the python-docx limitation that `paragraph.style = "Заголовок 1 Знак"`
    raises KeyError because the style is not registered in the document.
    """
    return SimpleNamespace(style=SimpleNamespace(name=name))


def test_classify_style_heading_en_ru() -> None:
    for name in [
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "heading 1",
        "Заголовок 1",
        "Заголовок 2",
        "Заголовок 1 Знак",
        "Заголовок Первого уровня (ГОСТ)",
    ]:
        assert classify_style(_paragraph_with_style_name(name)) == "heading", name


def test_classify_style_toc_en_ru() -> None:
    for name in [
        "toc 1",
        "toc 2",
        "toc 3",
        "TOC Heading",  # must be "toc" not "heading" — check order pinned here
        "содержание",
        "содержание основной",
        'Название "Содержание" (ГОСТ)',
    ]:
        assert classify_style(_paragraph_with_style_name(name)) == "toc", name


def test_classify_style_caption_en_ru() -> None:
    for name in [
        "Caption",
        "caption",
        "ПодписьРисунка",
        "Подпись рисунка (ГОСТ)",
        "Подпись таблицы (ГОСТ)",
    ]:
        assert classify_style(_paragraph_with_style_name(name)) == "caption", name


def test_classify_style_list_en_ru() -> None:
    for name in [
        "List Paragraph",
        "список",
        "Маркированный список (ГОСТ)",
        "Номерованный список (ГОСТ)",
    ]:
        assert classify_style(_paragraph_with_style_name(name)) == "list", name


def test_classify_style_body_negatives() -> None:
    for name in [
        "Normal",
        "Обычный",
        "Body Text",
        "header",
        "Footer",
        "",
        "mw-headline",  # contains "head" but not "heading" → must NOT match HEADING_STYLE_RE
    ]:
        assert classify_style(_paragraph_with_style_name(name)) == "body", name


def test_classify_style_handles_none_style() -> None:
    # paragraph.style is None — must NOT raise, must return "body"
    paragraph = SimpleNamespace(style=None)
    assert classify_style(paragraph) == "body"

    # Also covers the real-Document path where style.name might be None.
    document = Document()
    real_paragraph = document.add_paragraph("default Normal")
    assert classify_style(real_paragraph) == "body"


# ---- Phase 03 Wave 0 RED: heading signature extractor tests (D-02/D-03/D-04) ----


def _make_heading_paragraph(style="Heading 1"):
    document = Document()
    p = document.add_paragraph("Тест заголовка")
    p.style = style
    return p


def test_heading_signature_key_present() -> None:
    from src.rules.style_signatures import _extract_heading_format_signature
    p = _make_heading_paragraph()
    sig = _extract_heading_format_signature(p)
    assert isinstance(sig, dict)
    # 17 fields per D-02; assert presence of one from each category
    for key in ("font_name", "font_size", "bold", "italic", "underline", "color", "caps",
                "alignment", "first_line_indent_cm", "left_indent_cm", "right_indent_cm",
                "line_spacing", "space_before_pt", "space_after_pt",
                "keep_with_next", "keep_lines_together", "page_break_before", "widow_control"):
        assert key in sig, f"missing key {key} in heading_format_signature"
        assert isinstance(sig[key], dict), f"{key} entry must be dict"
        assert "value" in sig[key] and "source" in sig[key], f"{key} missing value/source"
        assert sig[key]["source"] in ("direct", "inherited", "unset"), f"{key} bad source: {sig[key]['source']}"


def test_heading_signature_direct_none_is_inherited() -> None:
    """D-03 Pass 1: when paragraph has NO direct overrides, every field source != 'direct'."""
    from src.rules.style_signatures import _extract_heading_format_signature
    p = _make_heading_paragraph()
    sig = _extract_heading_format_signature(p)
    for field, entry in sig.items():
        assert entry["source"] in ("inherited", "unset"), f"{field}: {entry}"


def test_heading_signature_direct_override_detected() -> None:
    """D-03 Pass 1 + D-04 source tagging: direct paragraph_format override → source='direct'."""
    from src.rules.style_signatures import _extract_heading_format_signature
    p = _make_heading_paragraph()
    p.paragraph_format.space_before = Pt(99.0)
    p.paragraph_format.first_line_indent = Cm(1.0)
    sig = _extract_heading_format_signature(p)
    assert sig["space_before_pt"]["source"] == "direct", sig["space_before_pt"]
    assert abs(sig["space_before_pt"]["value"] - 99.0) < 0.1, sig["space_before_pt"]
    assert sig["first_line_indent_cm"]["source"] == "direct", sig["first_line_indent_cm"]
    assert abs(sig["first_line_indent_cm"]["value"] - 1.0) < 0.05, sig["first_line_indent_cm"]


def test_heading_signature_cascade_walk() -> None:
    """D-03 Pass 2: walk style.base_style chain. Heading 1 inherits font.bold=True from style;
    paragraph has no direct override → source='inherited', value=True."""
    from src.rules.style_signatures import _extract_heading_format_signature
    p = _make_heading_paragraph(style="Heading 1")
    sig = _extract_heading_format_signature(p)
    # Heading 1's style font.bold is True in the default python-docx Document
    assert sig["bold"]["source"] in ("inherited", "unset"), sig["bold"]
    # If 'inherited', value must be the cascade-resolved value (bool)
    if sig["bold"]["source"] == "inherited":
        assert sig["bold"]["value"] is True, sig["bold"]
