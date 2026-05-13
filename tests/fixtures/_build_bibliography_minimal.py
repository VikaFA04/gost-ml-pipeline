"""One-shot fixture builder for tests/fixtures/bibliography_minimal.docx.

Run once:
    python tests/fixtures/_build_bibliography_minimal.py

Commits the resulting .docx as a binary fixture. Re-running is idempotent
in CONTENT (9 paragraphs, same styles, same mixed numId on entry 2), NOT
byte-identical (DOCX core.xml embeds creation timestamps).

Layout (D-14):
  1. "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" — Normal style (D-01 must override).
  2. "ТЕОРЕТИЧЕСКАЯ ЧАСТЬ" — Heading 1 (D-04 position+style detection).
  3. Entry — no numPr.
  4. Entry — numPr pointing at a legacy singleLevel abstract (numId=1) — D-06 coerce.
  5. Entry — no numPr.
  6. "ИССЛЕДОВАТЕЛЬСКИЙ РАЗДЕЛ" — Heading 1 (mixed naming, no number prefix).
  7. Entry — no numPr.
  8. Entry — no numPr.
  9. Entry — no numPr.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _next_abstract_num_id(numbering_root) -> str:
    """Allocate the next unused w:abstractNumId in this numbering_root.
    Mirrors src/rules/rule_engine._next_abstract_num_id but stays self-contained
    (the builder MUST NOT import rule_engine internals — Phase 2 is modifying
    them and the builder is the contract anchor for Plan 03's RED tests)."""
    existing = [
        int(a.get(qn("w:abstractNumId")))
        for a in numbering_root.findall(qn("w:abstractNum"))
        if a.get(qn("w:abstractNumId")) is not None
    ]
    return str((max(existing) if existing else -1) + 1)


def _next_num_id(numbering_root) -> str:
    """Allocate the next unused w:numId in this numbering_root.
    Mirrors src/rules/rule_engine._next_num_id; self-contained (same reason)."""
    existing = [
        int(n.get(qn("w:numId")))
        for n in numbering_root.findall(qn("w:num"))
        if n.get(qn("w:numId")) is not None
    ]
    return str((max(existing) if existing else 0) + 1)


def _inject_legacy_singlelevel_abstract(document) -> tuple[str, str]:
    """Append a singleLevel abstractNum (mimicking the pre-Phase-2 pattern) and
    a w:num pointing at it. Returns (abstract_num_id, num_id) so the caller
    knows which numId to attach to entry 2.

    Allocates IDs via _next_abstract_num_id / _next_num_id rather than
    hard-coding "0"/"1" — this is safe whether or not python-docx already
    created a numbering_part. Precondition: caller has not yet referenced
    the returned num_id from any paragraph.
    """
    numbering_part = document.part.numbering_part
    if numbering_part is None:
        # Force creation by adding a list-style paragraph then removing it.
        tmp = document.add_paragraph("__seed__")
        tmp.style = "List Number"
        numbering_part = document.part.numbering_part
        tmp._element.getparent().remove(tmp._element)
        assert numbering_part is not None, "python-docx failed to materialize numbering_part"
    numbering_root = numbering_part.element

    abstract_num_id = _next_abstract_num_id(numbering_root)
    num_id = _next_num_id(numbering_root)

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_num_id)
    mlt = OxmlElement("w:multiLevelType"); mlt.set(qn("w:val"), "singleLevel"); abstract_num.append(mlt)
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0")
    for tag, val in (("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "1.%1"), ("w:lvlJc", "left")):
        e = OxmlElement(tag); e.set(qn("w:val"), val); lvl.append(e)
    abstract_num.append(lvl)
    numbering_root.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), abstract_num_id); num.append(ref)
    numbering_root.append(num)

    return abstract_num_id, num_id


def _set_numPr(paragraph, num_id: str, ilvl: str = "0") -> None:
    """Set <w:numPr> on a paragraph with the given numId/ilvl."""
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl"); ilvl_el.set(qn("w:val"), ilvl); num_pr.append(ilvl_el)
    num_id_el = OxmlElement("w:numId"); num_id_el.set(qn("w:val"), num_id); num_pr.append(num_id_el)
    p_pr.append(num_pr)


def build(output_path: Path) -> None:
    document = Document()

    # Inject the legacy singleLevel abstract + a w:num BEFORE adding paragraphs
    # that reference it. _inject_legacy_singlelevel_abstract returns the allocated
    # (abstract_num_id, num_id) so we don't hard-code values that may collide with
    # IDs python-docx auto-created in numbering_part.
    legacy_abstract_id, legacy_num_id = _inject_legacy_singlelevel_abstract(document)

    # 1. Bibliography title — Normal style (D-01 must override SVM=body_text → bibliography_title).
    document.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")

    # 2. Subsection 1 — Heading 1, text NOT matching BIBLIOGRAPHY_SUBHEADING_RE (D-04 position+style).
    h1 = document.add_paragraph("ТЕОРЕТИЧЕСКАЯ ЧАСТЬ"); h1.style = "Heading 1"

    # 3. Entry 1 — no numPr.
    document.add_paragraph("Иванов И. И. Основы теории / И. И. Иванов. — Москва : Наука, 2020. — 240 с.")

    # 4. Entry 2 — stale numPr (numId=legacy_num_id, legacy singleLevel). D-06 must coerce.
    e2 = document.add_paragraph("Петров П. П. Введение в дискретную математику. — СПб. : Лань, 2019. — 320 с.")
    _set_numPr(e2, num_id=legacy_num_id, ilvl="0")

    # 5. Entry 3 — no numPr.
    document.add_paragraph("Сидоров С. С. Алгоритмы и структуры данных. — Москва : МЦНМО, 2018. — 188 с.")

    # 6. Subsection 2 — Heading 1, mixed naming (no number prefix), proves D-04 is style-based.
    h2 = document.add_paragraph("ИССЛЕДОВАТЕЛЬСКИЙ РАЗДЕЛ"); h2.style = "Heading 1"

    # 7-9. Entries — no numPr (D-06 fresh-allocate path).
    document.add_paragraph("Кузнецов А. А. Машинное обучение. — Москва : Питер, 2021. — 412 с.")
    document.add_paragraph("Морозов В. В. Нейронные сети. — Москва : ДМК Пресс, 2022. — 268 с.")
    document.add_paragraph("Лебедев Д. Д. Обработка естественного языка. — Москва : Эксмо, 2023. — 304 с.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


if __name__ == "__main__":
    build(Path("tests/fixtures/bibliography_minimal.docx"))
    print("wrote tests/fixtures/bibliography_minimal.docx")
