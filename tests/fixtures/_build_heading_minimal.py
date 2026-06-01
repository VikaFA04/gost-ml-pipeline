"""One-shot fixture builder for tests/fixtures/heading_minimal.docx.

Run once:
    python tests/fixtures/_build_heading_minimal.py

Commits the resulting .docx as a binary fixture. Re-running is idempotent
in CONTENT (4 paragraphs, same styles, same direct overrides per D-10), NOT
byte-identical (DOCX core.xml embeds creation timestamps).

Layout (D-10):
  1. Positive Heading 1 — target signature; all direct values None → zero fixes.
  2. Wrong-intervals heading — direct space_before_pt/space_after_pt override mismatch → autofix (D-06).
  3. Wrong-font-params heading — direct font_size/bold override mismatch → autofix (D-06).
  4. Inherited-mismatch heading — Heading 1 style-cascade value differs from profile → review (D-05).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def build(output_path: Path) -> None:
    document = Document()

    # 1. Positive — Heading 1, ALL direct overrides None
    p1 = document.add_paragraph("1 Основная часть")
    p1.style = "Heading 1"
    # No direct overrides set; all values inherited from Heading 1 style.

    # 2. Wrong-intervals — direct space_before/space_after override
    p2 = document.add_paragraph("2 Нарушение интервалов")
    p2.style = "Heading 1"
    p2.paragraph_format.space_before = Pt(99.0)  # direct override, wrong value
    p2.paragraph_format.space_after = Pt(99.0)

    # 3. Wrong-font-params — direct font_size/bold override on runs
    p3 = document.add_paragraph("3 Нарушение шрифта")
    p3.style = "Heading 1"
    for run in p3.runs:
        if run.text:
            run.font.size = Pt(9.0)   # direct override, wrong value
            run.bold = False           # direct override, wrong value

    # 4. Inherited-mismatch — style-cascade value differs from GOST profile
    # (Heading 1 style has space_after=28.35pt; profile expects 0pt)
    # No direct override → source="inherited" → D-05 review only
    p4 = document.add_paragraph("4 Унаследованное нарушение")
    p4.style = "Heading 1"
    # No direct override; inherited mismatch surfaces automatically.

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


if __name__ == "__main__":
    build(Path("tests/fixtures/heading_minimal.docx"))
    print("wrote tests/fixtures/heading_minimal.docx")
