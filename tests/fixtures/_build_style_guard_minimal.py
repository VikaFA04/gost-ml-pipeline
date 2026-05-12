"""One-shot fixture builder for tests/fixtures/style_guard_minimal.docx.

Run once: `python tests/fixtures/_build_style_guard_minimal.py`
Commits the resulting .docx as a binary fixture. Re-running is idempotent
in the sense that the same 5 styled paragraphs are produced (note: DOCX
core.xml embeds creation timestamps, so the file is NOT byte-identical
across rebuilds — see acceptance criteria below).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def build(output_path: Path) -> None:
    document = Document()

    p1 = document.add_paragraph("Глава 1. Введение")
    p1.style = "Heading 1"

    p2 = document.add_paragraph("Первый пункт перечисления")
    p2.style = "List Paragraph"

    p3 = document.add_paragraph("Глава 1 ............... 5")
    # python-docx default template ships "TOC Heading" (not "TOC 1");
    # classify_style maps "TOC Heading" to "toc" (pinned by
    # test_classify_style_toc_en_ru). RESEARCH §"Open Question 2" pins
    # check-order so that "TOC Heading" classifies as "toc" not "heading".
    p3.style = "TOC Heading"

    p4 = document.add_paragraph("Рисунок 1 — Схема")
    p4.style = "Caption"

    # Control: plain body paragraph (Normal style), long enough to look like body_text.
    document.add_paragraph(
        "Это обычный абзац основного текста, оставляемый стилем Normal "
        "в качестве контрольного образца — на нём существующий body_text "
        "путь должен срабатывать как раньше."
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


if __name__ == "__main__":
    build(Path("tests/fixtures/style_guard_minimal.docx"))
    print("wrote tests/fixtures/style_guard_minimal.docx")
