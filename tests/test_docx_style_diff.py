from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.evaluation.docx_style_diff import compare_docx_styles


def test_compare_docx_styles_reports_changed_paragraph_fields(tmp_path) -> None:
    positive_path = tmp_path / "positive.docx"
    positive_doc = Document()
    positive_doc.add_paragraph("Text")
    positive_doc.save(positive_path)

    candidate_path = tmp_path / "candidate.docx"
    candidate_doc = Document()
    candidate_paragraph = candidate_doc.add_paragraph("Text")
    candidate_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    candidate_doc.save(candidate_path)

    diff = compare_docx_styles(positive_path, candidate_path)

    assert diff.compared_paragraphs == 1
    assert diff.changed_paragraphs == 1
    assert diff.diff_rate == 1.0
    assert diff.field_mismatches["alignment"] == 1


def test_compare_docx_styles_identical_documents_have_zero_diff(tmp_path) -> None:
    positive_path = tmp_path / "positive.docx"
    document = Document()
    document.add_paragraph("Text")
    document.save(positive_path)

    diff = compare_docx_styles(positive_path, positive_path)

    assert diff.changed_paragraphs == 0
    assert diff.diff_rate == 0.0
    assert diff.field_mismatches == {}
