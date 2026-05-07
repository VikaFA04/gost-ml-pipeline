from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.rules.profile_loader import load_profile
from src.generate.docx_writer import generate_docx_from_predictions


def test_generate_docx_assigns_expected_styles(tmp_path) -> None:
    predictions_csv = tmp_path / "predictions.csv"
    output_docx = tmp_path / "output.docx"
    pd.DataFrame(
        [
            {
                "block_id": 1,
                "text": "Раздел 1",
                "kind": "paragraph",
                "predicted_label": "title_section",
            },
            {
                "block_id": 2,
                "text": "1. Пункт списка",
                "kind": "paragraph",
                "predicted_label": "list_item",
            },
            {
                "block_id": 3,
                "text": "Рисунок 1 — Схема",
                "kind": "paragraph",
                "predicted_label": "figure_caption",
            },
            {
                "block_id": 4,
                "text": "Обычный абзац",
                "kind": "paragraph",
                "predicted_label": "body_text",
            },
        ]
    ).to_csv(predictions_csv, index=False, encoding="utf-8-sig")

    generate_docx_from_predictions(predictions_csv, output_docx, profile_id="gost_7_32_2017")

    document = Document(output_docx)
    paragraphs = document.paragraphs
    styles = [paragraph.style.name for paragraph in paragraphs]

    assert styles == ["Heading 1", "List Paragraph", "Normal", "Normal"]
    assert paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert round(paragraphs[0].paragraph_format.left_indent.cm, 2) == 1.25
    assert round(paragraphs[0].paragraph_format.space_after.pt, 2) == 10.0
    assert paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert round(paragraphs[1].paragraph_format.left_indent.cm, 2) == 1.25
    assert round(paragraphs[1].paragraph_format.first_line_indent.cm, 2) == 0.0
    assert paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert round(paragraphs[2].paragraph_format.space_after.pt, 2) == 6.0
    assert round(paragraphs[3].paragraph_format.first_line_indent.cm, 2) == 1.25


def test_generate_docx_uses_custom_template_styles(tmp_path) -> None:
    template_path = tmp_path / "custom_template.docx"
    template = Document()
    for style_name in ["Custom H1", "Custom List", "Custom Normal", "Custom Caption"]:
        if style_name not in template.styles:
            template.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    template.save(template_path)

    profile = load_profile(profile_id="gost_7_32_2017")
    profile["document_rules"]["writer"] = {
        "template_path": str(template_path),
        "styles": {
            "title_section": {"style_name": "Custom H1"},
            "list_item": {"style_name": "Custom List"},
            "body_text": {"style_name": "Custom Normal"},
            "figure_caption": {"style_name": "Custom Caption"},
        },
    }
    profile_path = tmp_path / "custom_profile.json"
    profile_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")

    predictions_csv = tmp_path / "custom_predictions.csv"
    output_docx = tmp_path / "custom_output.docx"
    pd.DataFrame(
        [
            {
                "block_id": 1,
                "text": "Раздел 1",
                "kind": "paragraph",
                "predicted_label": "title_section",
            },
            {
                "block_id": 2,
                "text": "1. Пункт списка",
                "kind": "paragraph",
                "predicted_label": "list_item",
            },
            {
                "block_id": 3,
                "text": "Рисунок 1 — Схема",
                "kind": "paragraph",
                "predicted_label": "figure_caption",
            },
            {
                "block_id": 4,
                "text": "Обычный абзац",
                "kind": "paragraph",
                "predicted_label": "body_text",
            },
        ]
    ).to_csv(predictions_csv, index=False, encoding="utf-8-sig")

    generate_docx_from_predictions(predictions_csv, output_docx, profile_path=profile_path)

    document = Document(output_docx)
    styles = [paragraph.style.name for paragraph in document.paragraphs]

    assert styles == ["Custom H1", "Custom List", "Custom Caption", "Custom Normal"]


def test_generate_docx_writes_table_blocks(tmp_path) -> None:
    predictions_csv = tmp_path / "table_predictions.csv"
    output_docx = tmp_path / "table_output.docx"
    pd.DataFrame(
        [
            {
                "block_id": 1,
                "text": "A | B\n1 | 2",
                "kind": "table",
                "predicted_label": "body_text",
            },
        ]
    ).to_csv(predictions_csv, index=False, encoding="utf-8-sig")

    generate_docx_from_predictions(predictions_csv, output_docx)

    document = Document(output_docx)
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "A"
    assert document.tables[0].cell(0, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
