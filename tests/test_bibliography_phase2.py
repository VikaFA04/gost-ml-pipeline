"""Phase 02 RED tests — bibliography & list semantics.

D-01..D-15 coverage. Plans 02/03/04 implement; these MUST fail today.

File reference paths assume CWD == repository root (per pytest invocation).
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.evaluation.format_regression_audit import build_regression_predictions
from src.generate.inplace_formatter import audit_or_format_docx
from src.rules.rule_engine import (
    apply_bibliography_format,
    apply_rules_to_paragraph,
)
from src.rules.rule_loader import load_rules


# ---------------------- helpers ----------------------

def _row_data_body_text(text: str) -> dict:
    return {"text": text, "confidence_score": 0.99, "low_confidence": False}


def _bibliography_item_rows(report_df: pd.DataFrame) -> pd.DataFrame:
    """Filter the audit CSV to rows that the postprocess classified as bibliography_item.

    The audit CSV emits the active (postprocessed) label in the 'label' column.
    Phase 1 historical column names ('postprocessed_label', 'predicted_label')
    are kept as fallbacks for cross-version test runs.
    """
    for col in ("label", "postprocessed_label", "predicted_label"):
        if col in report_df.columns:
            return report_df[report_df[col] == "bibliography_item"].copy()
    raise KeyError(f"no label column found in report; columns={list(report_df.columns)!r}")


def _all_numIds_in_docx(docx_path: Path) -> list[str | None]:
    """For each bibliography_item-ish paragraph, return numId or None."""
    document = Document(str(docx_path))
    result: list[str | None] = []
    for p in document.paragraphs:
        try:
            p_pr = p._p.find(qn("w:pPr"))
            if p_pr is None:
                result.append(None); continue
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is None:
                result.append(None); continue
            num_id_el = num_pr.find(qn("w:numId"))
            result.append(num_id_el.get(qn("w:val")) if num_id_el is not None else None)
        except Exception:
            result.append(None)
    return result


# ============================================================
# D-09 — ambiguous-list review routing
# ============================================================

def test_ambiguous_list_marker_no_numId_routes_to_review() -> None:
    """D-09: body_text + numbered marker + no Word numPr + Normal style → review,
    explanation 'ambiguous_list_marker_no_numId', no fixes applied."""
    document = Document()
    paragraph = document.add_paragraph("1) Первый пункт без Word numbering, без подписки на список.")

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data=_row_data_body_text(paragraph.text),
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review", result
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert result["explanation"] == "ambiguous_list_marker_no_numId", result["explanation"]


# ============================================================
# D-10 — no marker + no numId + body_text stays body_text
# ============================================================

def test_long_body_text_without_marker_stays_body_text() -> None:
    """D-10: A long Normal-styled body paragraph without a list marker and
    without numPr must NOT receive list coercion. Phase 1 style guard +
    the absence of D-09 trigger leaves the existing body_text path to run.
    Concrete assertion: applied_fixes does NOT include 'numbering'; status
    is not 'review' with the D-09 explanation."""
    document = Document()
    long_text = (
        "Обычный абзац основного текста без маркера и без Word numbering, "
        "достаточно длинный, чтобы пройти MAX_FALLBACK_LIST_CHARS порог. " * 4
    )
    paragraph = document.add_paragraph(long_text)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data=_row_data_body_text(paragraph.text),
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    # Result may be None (no applicable body_text rules fire) or a non-D-09 dict.
    # The critical assertions: no numbering applied, no D-09 explanation.
    if result is not None:
        assert "numbering" not in result.get("applied_fixes", []), result
        assert result.get("explanation", "") != "ambiguous_list_marker_no_numId", result


# ============================================================
# D-05 — 2-level multilevel abstract emission
# ============================================================

def test_bibliography_multilevel_renders_section_dot_entry() -> None:
    """D-05: _create_bibliography_multilevel_abstract emits a multiLevelType
    multilevel abstract with TWO w:lvl children, level-1 lvlText='%1.%2.'."""
    from src.rules.rule_engine import _create_bibliography_multilevel_abstract  # RED via AttributeError today

    document = Document()
    # Trigger numbering_part creation.
    tmp = document.add_paragraph("__seed__"); tmp.style = "List Number"
    numbering_root = document.part.numbering_part.element

    abstract_num_id = _create_bibliography_multilevel_abstract(numbering_root)

    # Find the emitted abstract by id and inspect its structure.
    abstracts = numbering_root.findall(qn("w:abstractNum"))
    emitted = next((a for a in abstracts if a.get(qn("w:abstractNumId")) == abstract_num_id), None)
    assert emitted is not None, f"emitted abstractNumId={abstract_num_id} not found among {[a.get(qn('w:abstractNumId')) for a in abstracts]}"

    mlt = emitted.find(qn("w:multiLevelType"))
    assert mlt is not None and mlt.get(qn("w:val")) == "multilevel", mlt is not None and mlt.get(qn("w:val"))

    lvls = emitted.findall(qn("w:lvl"))
    assert len(lvls) == 2, f"expected 2 w:lvl children, got {len(lvls)}"
    ilvls = [l.get(qn("w:ilvl")) for l in lvls]
    assert "0" in ilvls and "1" in ilvls, ilvls

    lvl1 = next(l for l in lvls if l.get(qn("w:ilvl")) == "1")
    lvl1_text = lvl1.find(qn("w:lvlText"))
    assert lvl1_text is not None and lvl1_text.get(qn("w:val")) == "%1.%2.", (
        f"lvl-1 lvlText must be '%1.%2.', got {lvl1_text is not None and lvl1_text.get(qn('w:val'))!r}"
    )


def test_bibliography_num_with_section_override_carries_lvlOverride() -> None:
    """D-05 pitfall 2: each w:num MUST carry TWO w:lvlOverride children
    (ilvl=0 startOverride=section_index, ilvl=1 startOverride=1) so Word
    resets per-subsection counters correctly."""
    from src.rules.rule_engine import (
        _create_bibliography_multilevel_abstract,
        _create_bibliography_num_with_section_override,
    )  # RED via AttributeError today

    document = Document()
    tmp = document.add_paragraph("__seed__"); tmp.style = "List Number"
    numbering_root = document.part.numbering_part.element

    abstract_num_id = _create_bibliography_multilevel_abstract(numbering_root)
    num_id = _create_bibliography_num_with_section_override(numbering_root, abstract_num_id, section_index=2)

    nums = numbering_root.findall(qn("w:num"))
    emitted = next((n for n in nums if n.get(qn("w:numId")) == str(num_id)), None)
    assert emitted is not None

    overrides = emitted.findall(qn("w:lvlOverride"))
    assert len(overrides) == 2, f"expected 2 w:lvlOverride children, got {len(overrides)}"
    by_ilvl = {ov.get(qn("w:ilvl")): ov.find(qn("w:startOverride")).get(qn("w:val")) for ov in overrides}
    assert by_ilvl.get("0") == "2", f"lvlOverride ilvl=0 must startOverride=2, got {by_ilvl}"
    assert by_ilvl.get("1") == "1", f"lvlOverride ilvl=1 must startOverride=1, got {by_ilvl}"


# ============================================================
# D-06 — first-valid-numId coercion
# ============================================================

def test_bibliography_subsection_coerces_to_first_valid_numId(tmp_path) -> None:
    """D-06: bibliography_minimal.docx subsection 1 has mixed numIds (entry 2
    carries numId=1 pointing at legacy singleLevel). After apply_safe, at
    least one bibliography_item row's applied_fixes contains
    'numbering:coerced_to_numId=' (exact target numId varies)."""
    input_docx = Path("tests/fixtures/bibliography_minimal.docx")
    assert input_docx.exists(), "run tests/fixtures/_build_bibliography_minimal.py first"

    predictions_csv = tmp_path / "predictions.csv"
    report_csv = tmp_path / "report.csv"
    output_docx = tmp_path / "output.docx"
    build_regression_predictions(input_docx, predictions_csv)

    summary = audit_or_format_docx(
        input_docx=input_docx,
        predictions_csv=predictions_csv,
        report_csv=report_csv,
        output_docx=output_docx,
        apply_safe=True,
        profile_id="gost_7_32_2017",
    )
    assert summary["error"] == 0, summary

    report_df = pd.read_csv(report_csv, encoding="utf-8-sig")
    biblio_rows = _bibliography_item_rows(report_df)
    applied_fixes_concat = " | ".join(str(v) for v in biblio_rows.get("applied_fixes", []).tolist())
    assert "numbering:coerced_to_numId=" in applied_fixes_concat, (
        f"D-06 coercion tag not present. applied_fixes seen: {applied_fixes_concat!r}"
    )


# ============================================================
# D-07 — idempotent on re-run
# ============================================================

def test_bibliography_idempotent_on_rerun(tmp_path) -> None:
    """D-07: run apply_safe twice on bibliography_minimal.docx. Second run's
    summary['changed'] == 0 — numId already correct, no further fixes."""
    input_docx = Path("tests/fixtures/bibliography_minimal.docx")
    assert input_docx.exists()

    predictions_csv = tmp_path / "predictions.csv"
    report_csv_1 = tmp_path / "report_1.csv"
    output_docx_1 = tmp_path / "output_1.docx"
    build_regression_predictions(input_docx, predictions_csv)

    summary_1 = audit_or_format_docx(
        input_docx=input_docx, predictions_csv=predictions_csv,
        report_csv=report_csv_1, output_docx=output_docx_1,
        apply_safe=True, profile_id="gost_7_32_2017",
    )
    assert summary_1["error"] == 0

    # Re-feed the corrected DOCX.
    predictions_csv_2 = tmp_path / "predictions_2.csv"
    report_csv_2 = tmp_path / "report_2.csv"
    output_docx_2 = tmp_path / "output_2.docx"
    build_regression_predictions(output_docx_1, predictions_csv_2)

    summary_2 = audit_or_format_docx(
        input_docx=output_docx_1, predictions_csv=predictions_csv_2,
        report_csv=report_csv_2, output_docx=output_docx_2,
        apply_safe=True, profile_id="gost_7_32_2017",
    )
    assert summary_2["error"] == 0
    assert summary_2["changed"] == 0, (
        f"D-07: second run should produce no further changes, got changed={summary_2['changed']}. "
        f"Cache key {{id(numbering_root)}} likely leaking — switch to id(paragraph.part.document.part)."
    )


# ============================================================
# D-05 — bibliography_item paragraphs use ilvl=1 after fix
# ============================================================

def test_bibliography_apply_uses_ilvl_1(tmp_path) -> None:
    """D-05: After apply_safe, every bibliography_item paragraph's <w:numPr>
    has <w:ilvl w:val="1"/> (changed from "0" today)."""
    input_docx = Path("tests/fixtures/bibliography_minimal.docx")
    assert input_docx.exists()

    predictions_csv = tmp_path / "predictions.csv"
    report_csv = tmp_path / "report.csv"
    output_docx = tmp_path / "output.docx"
    build_regression_predictions(input_docx, predictions_csv)

    audit_or_format_docx(
        input_docx=input_docx, predictions_csv=predictions_csv,
        report_csv=report_csv, output_docx=output_docx,
        apply_safe=True, profile_id="gost_7_32_2017",
    )

    document = Document(str(output_docx))
    # 6 entries total (3 per subsection). Each must have numPr.ilvl == "1".
    biblio_paragraphs = [p for p in document.paragraphs if p.text.startswith(("Иванов", "Петров", "Сидоров", "Кузнецов", "Морозов", "Лебедев"))]
    assert len(biblio_paragraphs) == 6, f"expected 6 bibliography entries, found {len(biblio_paragraphs)}"
    for p in biblio_paragraphs:
        p_pr = p._p.find(qn("w:pPr"))
        assert p_pr is not None, f"missing pPr on entry {p.text[:30]!r}"
        num_pr = p_pr.find(qn("w:numPr"))
        assert num_pr is not None, f"missing numPr on entry {p.text[:30]!r}"
        ilvl = num_pr.find(qn("w:ilvl"))
        assert ilvl is not None and ilvl.get(qn("w:val")) == "1", (
            f"entry {p.text[:30]!r} ilvl={ilvl is not None and ilvl.get(qn('w:val'))!r}, expected '1'"
        )


# ============================================================
# D-13 — bibliography_format skips alignment when profile omits field
# ============================================================

def test_bibliography_format_skips_alignment_when_profile_omits() -> None:
    """D-13: apply_bibliography_format with config={'style_name':'List Number'}
    (no alignment / indent fields) must NOT write alignment. The paragraph's
    alignment stays None after the call."""
    document = Document()
    paragraph = document.add_paragraph("Иванов И. И. Тест.")

    config = {"style_name": "List Number"}  # NO alignment, NO indents
    applied = apply_bibliography_format(paragraph, config, section_index=1)

    # Strict assertion: alignment must not appear in applied_fixes — profile didn't ask for it.
    assert "alignment" not in applied, f"applied_fixes contains 'alignment' but profile didn't carry it: {applied}"
    assert paragraph.alignment is None, (
        f"paragraph.alignment={paragraph.alignment!r} — apply_bibliography_format wrote a direct alignment "
        "even though profile config did not carry the field. D-13 violated."
    )


# ============================================================
# D-14 — bibliography_minimal.docx single numId per subsection
# ============================================================

def test_bibliography_minimal_docx_single_numId_per_subsection(tmp_path) -> None:
    """D-14 hand-crafted: after apply_safe, all 3 entries in subsection 1
    share one numId; all 3 entries in subsection 2 share another numId;
    the two numIds differ (per-subsection scope per D-03 default)."""
    input_docx = Path("tests/fixtures/bibliography_minimal.docx")
    assert input_docx.exists()

    predictions_csv = tmp_path / "predictions.csv"
    report_csv = tmp_path / "report.csv"
    output_docx = tmp_path / "output.docx"
    build_regression_predictions(input_docx, predictions_csv)

    summary = audit_or_format_docx(
        input_docx=input_docx, predictions_csv=predictions_csv,
        report_csv=report_csv, output_docx=output_docx,
        apply_safe=True, profile_id="gost_7_32_2017",
    )
    assert summary["error"] == 0

    all_num_ids = _all_numIds_in_docx(output_docx)
    # paragraphs 2,3,4 = subsection 1 entries; 6,7,8 = subsection 2 entries.
    sub1_num_ids = [all_num_ids[2], all_num_ids[3], all_num_ids[4]]
    sub2_num_ids = [all_num_ids[6], all_num_ids[7], all_num_ids[8]]

    assert len(set(sub1_num_ids)) == 1 and sub1_num_ids[0] is not None, (
        f"subsection 1 entries must share one numId, got {sub1_num_ids}"
    )
    assert len(set(sub2_num_ids)) == 1 and sub2_num_ids[0] is not None, (
        f"subsection 2 entries must share one numId, got {sub2_num_ids}"
    )
    assert sub1_num_ids[0] != sub2_num_ids[0], (
        f"per_section scope: subsection 1 and subsection 2 must have DIFFERENT numIds, got both = {sub1_num_ids[0]}"
    )


# ============================================================
# D-14 — negative integration on real DOCX
# ============================================================

def test_negative_4_bibliography_single_numId(tmp_path) -> None:
    """D-14 negative integration: negative_examples/4_formatted_20260413_185420.docx
    has Heading 2 subsections and existing numId=16. After apply_safe, all
    bibliography_item rows share one numId per subsection; applied_fixes for
    at least one row includes 'numbering'."""
    input_docx = Path("negative_examples/4_formatted_20260413_185420.docx")
    if not input_docx.exists():
        pytest.skip(f"fixture {input_docx} not present in this environment")

    predictions_csv = tmp_path / "predictions.csv"
    report_csv = tmp_path / "report.csv"
    output_docx = tmp_path / "output.docx"
    build_regression_predictions(input_docx, predictions_csv)

    summary = audit_or_format_docx(
        input_docx=input_docx, predictions_csv=predictions_csv,
        report_csv=report_csv, output_docx=output_docx,
        apply_safe=True, profile_id="gost_7_32_2017",
    )
    assert summary["error"] == 0

    report_df = pd.read_csv(report_csv, encoding="utf-8-sig")
    biblio_rows = _bibliography_item_rows(report_df)
    assert not biblio_rows.empty, "expected ≥1 bibliography_item row in 4_formatted_20260413_185420.docx"

    applied_concat = " | ".join(str(v) for v in biblio_rows.get("applied_fixes", []).tolist())
    assert "numbering" in applied_concat, (
        f"expected at least one bibliography_item row's applied_fixes to include 'numbering', got {applied_concat!r}"
    )


def test_negative_3_bibliography_coerces_mixed_numIds(tmp_path) -> None:
    """D-06 + D-14: negative_examples/3_formatted_20260413_194927.docx carries
    mixed numIds (some None, some numId=1). Coercion fires → at least one row's
    applied_fixes contains 'numbering:coerced_to_numId='.

    If the document does not in fact exhibit mixed numIds on bibliography_item
    rows (researcher's claim — verify in implementation), this test should
    still pass because the fresh-allocate path also tags 'numbering'. The
    strict coercion-tag assertion lives in test_bibliography_subsection_coerces_to_first_valid_numId
    on bibliography_minimal.docx which is known-mixed.
    """
    input_docx = Path("negative_examples/3_formatted_20260413_194927.docx")
    if not input_docx.exists():
        pytest.skip(f"fixture {input_docx} not present in this environment")

    predictions_csv = tmp_path / "predictions.csv"
    report_csv = tmp_path / "report.csv"
    output_docx = tmp_path / "output.docx"
    build_regression_predictions(input_docx, predictions_csv)

    summary = audit_or_format_docx(
        input_docx=input_docx, predictions_csv=predictions_csv,
        report_csv=report_csv, output_docx=output_docx,
        apply_safe=True, profile_id="gost_7_32_2017",
    )
    assert summary["error"] == 0

    report_df = pd.read_csv(report_csv, encoding="utf-8-sig")
    biblio_rows = _bibliography_item_rows(report_df)
    assert not biblio_rows.empty, "expected ≥1 bibliography_item row in 3_formatted_20260413_194927.docx"
    applied_concat = " | ".join(str(v) for v in biblio_rows.get("applied_fixes", []).tolist())
    assert "numbering" in applied_concat, (
        f"expected 'numbering' in applied_fixes for at least one row, got {applied_concat!r}"
    )
