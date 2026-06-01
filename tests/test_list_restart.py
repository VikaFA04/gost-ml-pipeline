"""Defect 2 — every list restarts numbering at 1.

A new list (a separate list_instance) must NOT continue the previous list's
numbering. apply_list_numbering allocates one fresh numId per (document,
list_instance) and stamps it with a level startOverride=1 so Word restarts
the counter for each list.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

from src.rules.rule_engine import apply_list_numbering, apply_rules_to_paragraph
from src.rules.rule_loader import load_rules

FIXTURE = Path("tests/fixtures/bibliography_minimal.docx")


def _num_id_of(paragraph) -> str | None:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("w:numId"))
    return num_id.get(qn("w:val")) if num_id is not None else None


def _start_override_of(numbering_root, num_id: str) -> str | None:
    for num in numbering_root.findall(qn("w:num")):
        if num.get(qn("w:numId")) == str(num_id):
            override = num.find(qn("w:lvlOverride"))
            if override is None:
                return None
            start = override.find(qn("w:startOverride"))
            return start.get(qn("w:val")) if start is not None else None
    return None


def test_each_list_instance_restarts_at_1() -> None:
    document = Document(str(FIXTURE))
    numbering_root = document.part.numbering_part.element

    p1 = document.add_paragraph("первый пункт первого списка")
    p2 = document.add_paragraph("второй пункт первого списка")
    p3 = document.add_paragraph("первый пункт второго списка")

    apply_list_numbering(p1, "numbered", list_instance=1, list_level=0)
    apply_list_numbering(p2, "numbered", list_instance=1, list_level=0)
    apply_list_numbering(p3, "numbered", list_instance=2, list_level=0)

    num1, num2, num3 = _num_id_of(p1), _num_id_of(p2), _num_id_of(p3)

    # Same list instance shares one numId; a separate instance gets its own.
    assert num1 is not None and num1 == num2
    assert num3 is not None and num3 != num1

    # Each list's numId carries a startOverride=1 so Word restarts at 1.
    assert _start_override_of(numbering_root, num1) == "1"
    assert _start_override_of(numbering_root, num3) == "1"


def test_row_data_list_instance_drives_restart_through_rule_engine() -> None:
    """End-to-end wiring: a list_item row carrying list_instance flows through
    apply_rules_to_paragraph → apply_list_format → apply_list_numbering and the
    paragraph ends up with a per-instance numId that restarts at 1."""
    document = Document(str(FIXTURE))
    numbering_root = document.part.numbering_part.element

    paragraph = document.add_paragraph("первый пункт списка задач")
    # Pre-seed valid Word numbering so the block is structurally a real list
    # (safe_to_autofix), and give it a non-GOST layout so the layout rule fires.
    apply_list_numbering(paragraph, "numbered")
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.first_line_indent = Cm(0.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "первый пункт списка задач",
            "list_type": "numbered",
            "list_level": 0,
            "list_instance": 1,
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    num_id = _num_id_of(paragraph)
    assert num_id is not None
    assert _start_override_of(numbering_root, num_id) == "1"


def test_get_list_num_id_revalidates_stale_cache_entry() -> None:
    """Defect-2 review #1: the per-(doc,instance) numId cache is keyed by
    id(document.part), which CPython reuses after GC. A stale entry from a
    freed document must NOT leak a numId that does not exist in the current
    document — the allocator must re-validate against the live numbering."""
    from src.rules import rule_engine

    rule_engine._LIST_NUM_IDS.clear()
    document = Document(str(FIXTURE))
    numbering_root = document.part.numbering_part.element
    p = document.add_paragraph("пункт списка")

    # Simulate id() reuse: a numId cached under this doc's key but absent here.
    rule_engine._LIST_NUM_IDS[(rule_engine._document_cache_key(p), 1)] = 99999
    try:
        rule_engine.apply_list_numbering(p, "numbered", list_instance=1, list_level=0)
        nid = _num_id_of(p)
        assert nid != "99999", "leaked stale cross-document numId"
        assert any(
            n.get(qn("w:numId")) == nid for n in numbering_root.findall(qn("w:num"))
        ), "assigned numId must exist in this document's numbering"
    finally:
        rule_engine._LIST_NUM_IDS.clear()


def test_restart_preserves_existing_nested_numbering() -> None:
    """Defect-2 review #2: forcing a per-instance restart must not flatten a
    paragraph that already carries valid nested numbering (ilvl > 0)."""
    from src.rules import rule_engine

    rule_engine._LIST_NUM_IDS.clear()
    document = Document(str(FIXTURE))
    p = document.add_paragraph("вложенный пункт")
    apply_list_numbering(p, "numbered")  # valid numPr at ilvl=0
    numpr = p._p.pPr.find(qn("w:numPr"))
    numpr.find(qn("w:ilvl")).set(qn("w:val"), "2")  # source is nested at level 2
    num_before = numpr.find(qn("w:numId")).get(qn("w:val"))

    apply_list_numbering(p, "numbered", list_instance=1, list_level=0)

    numpr = p._p.pPr.find(qn("w:numPr"))
    assert numpr.find(qn("w:ilvl")).get(qn("w:val")) == "2", "flattened nested level"
    assert numpr.find(qn("w:numId")).get(qn("w:val")) == num_before, "clobbered valid nested numId"
    rule_engine._LIST_NUM_IDS.clear()
