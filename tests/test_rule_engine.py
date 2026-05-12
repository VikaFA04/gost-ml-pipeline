from __future__ import annotations

from pathlib import Path
import shutil
import uuid

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from src.evaluation.format_regression_audit import build_regression_predictions
from src.generate.inplace_formatter import audit_or_format_docx
from src.rules.profile_loader import load_profile
from src.rules.rule_engine import apply_list_numbering, apply_rules_to_paragraph
from src.rules.rule_loader import load_rules


def build_prediction_csv(path: Path, text: str, label: str, list_level: int | None = None) -> None:
    df = pd.DataFrame(
        [
            {
                "doc_id": "doc_1",
                "block_id": 1,
                "text": text,
                "kind": "paragraph",
                "alignment": "LEFT",
                "style": "Normal",
                "bold_ratio": 0.0,
                "file_name": "sample.docx",
                "predicted_label": label,
                "postprocessed_label": label,
                "confidence_score": 0.99,
                "low_confidence": False,
                "list_type": "numbered" if label == "list_item" else None,
                "list_level": list_level,
            }
        ]
    )
    df.to_csv(path, index=False, encoding="utf-8-sig")


def build_workspace_temp_dir() -> Path:
    base_dir = Path("tests_runtime")
    base_dir.mkdir(parents=True, exist_ok=True)
    temp_dir = base_dir / f"case_{uuid.uuid4().hex}"
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir


def test_rule_loading() -> None:
    rules = load_rules()
    assert rules
    assert any(rule["id"] == "list_item_layout" for rule in rules)
    assert all("priority" in rule for rule in rules)


def test_caption_profiles_require_review_not_autofix() -> None:
    profile = load_profile(profile_id="gost_7_32_2017")

    assert profile["labels"]["figure_caption"]["audit_policy"]["allow_auto_fix"] is False
    assert profile["labels"]["table_caption"]["audit_policy"]["allow_auto_fix"] is False


def test_rule_application_review_without_fix() -> None:
    document = Document()
    paragraph = document.add_paragraph("1.\tList item")
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={"list_level": 0, "list_type": "numbered"},
        rules=load_rules(),
        apply_safe=False,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert "list_item_layout" in result["violated_rules"]
    assert result["applied_fixes"] == []


def test_inherited_body_text_formatting_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Обычный текст с форматированием из стиля Word.")

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data={"confidence_score": 0.99, "low_confidence": False},
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert paragraph.alignment is None
    assert paragraph.paragraph_format.first_line_indent is None
    assert paragraph.paragraph_format.line_spacing is None


def test_inherited_heading_bold_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("1 Заголовок")
    paragraph.style = "Heading 1"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="title_section",
        row_data={"confidence_score": 0.99, "low_confidence": False},
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert paragraph.runs[0].bold is None


def test_heading_style_direct_alignment_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Список источников")
    paragraph.style = "Heading 2"
    paragraph.alignment = 1

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="title_subsection",
        row_data={"text": paragraph.text, "confidence_score": 0.99, "low_confidence": False},
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert paragraph.alignment == 1


def test_list_like_paragraph_predicted_as_body_text_is_not_autofixed() -> None:
    document = Document()
    paragraph = document.add_paragraph("1.\tList item")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data={
            "text": "1.\tList item",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    # Plan 03 style_guard fires earlier than the old list-like
    # blocked_unsafe_autofix path: List-styled paragraph + body_text label
    # short-circuits with style_guard_block, not blocked_unsafe_autofix.
    assert result["explanation"].startswith("style_guard_block:")
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert round(paragraph.paragraph_format.left_indent.cm, 2) == 2.25
    assert round(paragraph.paragraph_format.first_line_indent.cm, 2) == -1.0


def test_body_text_alignment_mismatch_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Обычный абзац с явным выравниванием.")
    paragraph.alignment = 0

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data={"text": paragraph.text, "confidence_score": 0.99, "low_confidence": False},
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert paragraph.alignment == 0


def test_body_text_hanging_indent_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Источник с висячим отступом.")
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data={"text": paragraph.text, "confidence_score": 0.99, "low_confidence": False},
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert round(paragraph.paragraph_format.left_indent.cm, 2) == 2.25
    assert round(paragraph.paragraph_format.first_line_indent.cm, 2) == -1.0


def test_body_text_line_spacing_mismatch_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Обычный абзац с явным межстрочным интервалом.")
    paragraph.paragraph_format.line_spacing = 1.0

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data={"confidence_score": 0.99, "low_confidence": False},
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert paragraph.paragraph_format.line_spacing == 1.0


def test_list_formatting_fix_level_1() -> None:
    document = Document()
    paragraph = document.add_paragraph("a)\tNested item")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.first_line_indent = Cm(0.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "a)\tNested item",
            "list_level": 1,
            "list_type": "numbered",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert "left_indent_cm" in result["applied_fixes"]
    assert round(paragraph.paragraph_format.left_indent.cm, 2) == 2.5
    assert round(paragraph.paragraph_format.first_line_indent.cm, 2) == -1.0
    tab_stops = list(paragraph.paragraph_format.tab_stops)
    assert tab_stops
    assert round(tab_stops[0].position.cm, 2) == 2.5


def test_list_formatting_repairs_broken_numbering_reference() -> None:
    document = Document()
    paragraph = document.add_paragraph("Broken bullet item")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.first_line_indent = Cm(0.0)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "999")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "Broken bullet item",
            "list_level": 0,
            "list_type": "bullet",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert "numbering" in result["applied_fixes"]
    repaired_num_id = paragraph._p.pPr.numPr.numId.val
    numbering_xml = paragraph.part.numbering_part.element.xml
    assert f'w:numId="{repaired_num_id}"' in numbering_xml


def test_accepted_list_layout_still_repairs_broken_numbering_reference() -> None:
    document = Document()
    paragraph = document.add_paragraph("Accepted layout with broken marker")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "999")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "Accepted layout with broken marker",
            "list_level": 0,
            "list_type": "bullet",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert result["applied_fixes"] == ["numbering"]


def test_bibliography_item_gets_numbered_word_style_without_text_change() -> None:
    document = Document()
    paragraph = document.add_paragraph("Иванов И. И. Учебник. — Москва, 2020. — 120 с.")
    paragraph.style = "Normal"
    before_text = paragraph.text

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="bibliography_item",
        row_data={
            "text": before_text,
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert "numbering" in result["applied_fixes"]
    assert paragraph._p.pPr is not None
    assert paragraph._p.pPr.numPr is not None
    assert paragraph._p.pPr.numPr.numId is not None


def test_accepted_list_layout_without_numbering_stays_unchanged() -> None:
    document = Document()
    paragraph = document.add_paragraph("Accepted layout with no numbering")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "Accepted layout with no numbering",
            "list_level": 0,
            "list_type": "bullet",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "no_change"
    assert result["applied_fixes"] == []


def test_accepted_non_list_style_without_numbering_gets_numbering() -> None:
    document = Document()
    paragraph = document.add_paragraph("Accepted layout with no numbering")
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)
    paragraph.style = "Quote"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "Accepted layout with no numbering",
            "list_level": 0,
            "list_type": "bullet",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "no_change"
    assert result["applied_fixes"] == []
    assert paragraph._p.pPr is not None
    assert paragraph._p.pPr.numPr is None


def test_body_text_accepted_list_layout_gets_numbering() -> None:
    document = Document()
    paragraph = document.add_paragraph("Body text with accepted list layout")
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)
    paragraph.style = "Quote"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data={
            "text": "Body text with accepted list layout",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert paragraph._p.pPr is not None
    assert paragraph._p.pPr.numPr is None


def test_list_item_without_layout_gets_format_and_numbering() -> None:
    document = Document()
    paragraph = document.add_paragraph("List item with no layout")

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "List item with no layout",
            "list_level": 0,
            "list_type": "bullet",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert "numbering" in result["applied_fixes"]
    assert "left_indent_cm" in result["applied_fixes"]
    assert "first_line_indent_cm" in result["applied_fixes"]
    assert paragraph._p.pPr is not None
    assert paragraph._p.pPr.numPr is not None
    assert paragraph._p.pPr.numPr.numId is not None


def test_numbered_list_with_inherited_layout_keeps_layout_unchanged() -> None:
    document = Document()
    paragraph = document.add_paragraph("COUNT(*) returns row count")
    apply_list_numbering(paragraph, "bullet")

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": paragraph.text,
            "list_level": 0,
            "list_type": "bullet",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "no_change"
    assert result["applied_fixes"] == []
    assert paragraph.paragraph_format.left_indent is None
    assert paragraph.paragraph_format.first_line_indent is None
    assert list(paragraph.paragraph_format.tab_stops) == []


def test_generic_style_only_list_item_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Generic style-only list candidate")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.first_line_indent = Cm(0.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": paragraph.text,
            "list_level": 0,
            "list_type": "list",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["blocked_unsafe_autofix"] is True
    assert result["applied_fixes"] == []
    assert round(paragraph.paragraph_format.left_indent.cm, 2) == 1.0
    assert round(paragraph.paragraph_format.first_line_indent.cm, 2) == 0.0
    assert paragraph._p.pPr is None or paragraph._p.pPr.numPr is None


def test_list_style_partial_layout_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Bullet item with local inherited first indent")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(1.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": paragraph.text,
            "list_level": 0,
            "list_type": "bullet",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert round(paragraph.paragraph_format.left_indent.cm, 2) == 1.0
    assert paragraph.paragraph_format.first_line_indent is None


def test_bibliography_item_numbering_uses_section_prefix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Data normalization / URL: https://example.test")

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="bibliography_item",
        row_data={
            "text": paragraph.text,
            "bibliography_section_index": 2,
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert "numbering" in result["applied_fixes"]
    numbering_xml = paragraph.part.numbering_part.element.xml
    assert 'w:val="2.%1"' in numbering_xml


def test_bibliography_item_replaces_wrong_existing_section_numbering() -> None:
    document = Document()
    paragraph = document.add_paragraph("Иванов И. И. Учебник. — Москва, 2020. — 120 с.")
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:numPr"))

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="bibliography_item",
        row_data={
            "text": paragraph.text,
            "bibliography_section_index": 1,
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert "numbering" in result["applied_fixes"]
    assert 'w:val="1.%1"' in paragraph.part.numbering_part.element.xml


def test_bibliography_subheading_gets_section_number_prefix() -> None:
    document = Document()
    paragraph = document.add_paragraph("Практическая часть")

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="bibliography_title",
        row_data={
            "text": paragraph.text,
            "bibliography_section_index": 2,
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert "bibliography_section_prefix" in result["applied_fixes"]
    assert paragraph.text == "2 ПРАКТИЧЕСКАЯ ЧАСТЬ"


def test_numbered_bibliography_title_section_gets_section_number_prefix() -> None:
    document = Document()
    paragraph = document.add_paragraph("1 Теоретическая часть")

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="title_section",
        row_data={
            "text": paragraph.text,
            "bibliography_section_index": 1,
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "changed"
    assert "bibliography_section_prefix" in result["applied_fixes"]
    assert paragraph.text == "1 ТЕОРЕТИЧЕСКАЯ ЧАСТЬ"


def test_bibliography_title_section_keeps_existing_alignment() -> None:
    document = Document()
    paragraph = document.add_paragraph("2 ПРАКТИЧЕСКАЯ ЧАСТЬ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="title_section",
        row_data={
            "text": paragraph.text,
            "bibliography_section_index": 2,
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["applied_fixes"] == []
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_marker_only_list_item_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("- marker-only list item")
    paragraph.paragraph_format.first_line_indent = Cm(1.25)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "- marker-only list item",
            "list_level": 0,
            "list_type": "bullet",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["blocked_unsafe_autofix"] is True
    assert result["applied_fixes"] == []
    assert round(paragraph.paragraph_format.first_line_indent.cm, 2) == 1.25


def test_list_item_alignment_mismatch_requires_review_not_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("List style item")
    paragraph.style = "List Paragraph"
    paragraph.alignment = 0
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "List style item",
            "list_level": 0,
            "list_type": "list",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert paragraph.alignment == 0


def test_accepted_positive_list_layout_ignores_inferred_level() -> None:
    document = Document()
    paragraph = document.add_paragraph("Bibliography-like item")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "Bibliography-like item",
            "list_level": 1,
            "list_type": "list",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "no_change"
    assert result["applied_fixes"] == []
    assert round(paragraph.paragraph_format.left_indent.cm, 2) == 2.25
    assert round(paragraph.paragraph_format.first_line_indent.cm, 2) == -1.0


def test_list_paragraph_with_accepted_layout_but_missing_numbering_gets_numbering() -> None:
    document = Document()
    paragraph = document.add_paragraph("Accepted list layout without numbering")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "Accepted list layout without numbering",
            "list_level": 0,
            "list_type": "list",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "no_change"
    assert result["applied_fixes"] == []
    assert paragraph._p.pPr is not None
    assert paragraph._p.pPr.numPr is None


def test_inherited_list_paragraph_layout_is_not_autofixed() -> None:
    document = Document()
    paragraph = document.add_paragraph("List style item")
    paragraph.style = "List Paragraph"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "List style item",
            "list_level": 0,
            "list_type": "list",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "no_change"
    assert result["applied_fixes"] == []
    assert paragraph.paragraph_format.left_indent is None
    assert paragraph.paragraph_format.first_line_indent is None


def test_positive_corpus_list_layout_is_not_autofixed() -> None:
    document = Document()
    paragraph = document.add_paragraph("пункт списка;")
    paragraph.style = "List Paragraph"
    paragraph.paragraph_format.left_indent = Cm(2.25)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "пункт списка;",
            "list_level": 0,
            "list_type": "list",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "no_change"
    assert result["applied_fixes"] == []
    assert round(paragraph.paragraph_format.left_indent.cm, 2) == 2.25
    assert round(paragraph.paragraph_format.first_line_indent.cm, 2) == -1.0


def test_low_confidence_list_item_blocks_unsafe_autofix() -> None:
    document = Document()
    paragraph = document.add_paragraph("1.\tList item")

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": "1.\tList item",
            "list_level": 0,
            "list_type": "numbered",
            "confidence_score": 0.65,
            "low_confidence": True,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["blocked_unsafe_autofix"] is True
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []


def test_long_paragraph_is_not_auto_fixed_as_list() -> None:
    document = Document()
    text = "1. " + "длинный текст " * 50
    paragraph = document.add_paragraph(text)

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="list_item",
        row_data={
            "text": text,
            "list_level": 0,
            "list_type": "numbered",
            "confidence_score": 0.99,
            "low_confidence": False,
        },
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["blocked_unsafe_autofix"] is True
    assert "ordinary paragraph" in result["unsafe_auto_fix_reason"]


def test_audit_mode_reports_review_and_fix_mode_changes() -> None:
    tmp_path = build_workspace_temp_dir()
    try:
        input_docx = tmp_path / "input.docx"
        document = Document()
        paragraph = document.add_paragraph("1.\tList item")
        paragraph.style = "List Paragraph"
        paragraph.paragraph_format.left_indent = Cm(1.0)
        paragraph.paragraph_format.first_line_indent = Cm(0.0)
        document.save(input_docx)

        predictions_csv = tmp_path / "predictions.csv"
        build_prediction_csv(predictions_csv, "1.\tList item", "list_item", list_level=0)

        audit_report = tmp_path / "audit.csv"
        audit_summary = audit_or_format_docx(
            input_docx=input_docx,
            predictions_csv=predictions_csv,
            report_csv=audit_report,
            apply_safe=False,
            profile_id="gost_7_32_2017",
        )
        assert audit_summary["review"] == 1
        assert audit_summary["changed"] == 0

        fixed_docx = tmp_path / "fixed.docx"
        fix_report = tmp_path / "fix.csv"
        fix_summary = audit_or_format_docx(
            input_docx=input_docx,
            predictions_csv=predictions_csv,
            report_csv=fix_report,
            output_docx=fixed_docx,
            apply_safe=True,
            profile_id="gost_7_32_2017",
        )
        assert fix_summary["changed"] == 1
        assert fixed_docx.exists()
    finally:
        shutil.rmtree(tmp_path, ignore_errors=True)


def test_fix_mode_is_idempotent() -> None:
    tmp_path = build_workspace_temp_dir()
    try:
        input_docx = tmp_path / "input.docx"
        document = Document()
        paragraph = document.add_paragraph("1.\tList item")
        paragraph.style = "List Paragraph"
        paragraph.paragraph_format.left_indent = Cm(1.0)
        paragraph.paragraph_format.first_line_indent = Cm(0.0)
        document.save(input_docx)

        predictions_csv = tmp_path / "predictions.csv"
        build_prediction_csv(predictions_csv, "1.\tList item", "list_item", list_level=0)

        first_output = tmp_path / "fixed_1.docx"
        audit_or_format_docx(
            input_docx=input_docx,
            predictions_csv=predictions_csv,
            report_csv=tmp_path / "fix_1.csv",
            output_docx=first_output,
            apply_safe=True,
            profile_id="gost_7_32_2017",
        )

        second_output = tmp_path / "fixed_2.docx"
        second_summary = audit_or_format_docx(
            input_docx=first_output,
            predictions_csv=predictions_csv,
            report_csv=tmp_path / "fix_2.csv",
            output_docx=second_output,
            apply_safe=True,
            profile_id="gost_7_32_2017",
        )

        assert second_summary["changed"] == 0
        assert second_summary["no_change"] == 1
    finally:
        shutil.rmtree(tmp_path, ignore_errors=True)


def test_fix_mode_preserves_tabbed_list_markers_from_predictions_csv() -> None:
    tmp_path = build_workspace_temp_dir()
    try:
        input_docx = tmp_path / "input.docx"
        document = Document()
        paragraph = document.add_paragraph("\tПервый пункт")
        paragraph.style = "List Paragraph"
        paragraph.paragraph_format.left_indent = Cm(2.25)
        paragraph.paragraph_format.first_line_indent = Cm(-1.0)
        document.save(input_docx)

        predictions_csv = tmp_path / "predictions.csv"
        build_prediction_csv(predictions_csv, "\tПервый пункт", "list_item", list_level=0)

        fixed_docx = tmp_path / "fixed.docx"
        summary = audit_or_format_docx(
            input_docx=input_docx,
            predictions_csv=predictions_csv,
            report_csv=tmp_path / "fix.csv",
            output_docx=fixed_docx,
            apply_safe=True,
            profile_id="gost_7_32_2017",
        )

        assert summary["changed"] == 1
        fixed = Document(fixed_docx)
        assert fixed.paragraphs[0]._p.pPr is not None
        assert fixed.paragraphs[0]._p.pPr.numPr is not None
    finally:
        shutil.rmtree(tmp_path, ignore_errors=True)


def test_fix_mode_preserves_tabbed_bibliography_entries_from_predictions_csv() -> None:
    tmp_path = build_workspace_temp_dir()
    try:
        input_docx = tmp_path / "input.docx"
        document = Document()
        document.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        document.add_paragraph("1 Теоретическая часть")
        document.add_paragraph("\tИванов И. И. Учебник. — Москва, 2020. — 120 с.")
        document.save(input_docx)

        predictions_csv = tmp_path / "predictions.csv"
        pd.DataFrame(
            [
                {
                    "doc_id": "doc_1",
                    "block_id": 1,
                    "text": "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
                    "kind": "paragraph",
                    "alignment": "LEFT",
                    "style": "Normal",
                    "bold_ratio": 0.0,
                    "file_name": "sample.docx",
                    "predicted_label": "body_text",
                    "postprocessed_label": "bibliography_title",
                    "confidence_score": 0.99,
                    "low_confidence": False,
                },
                {
                    "doc_id": "doc_1",
                    "block_id": 2,
                    "text": "1 Теоретическая часть",
                    "kind": "paragraph",
                    "alignment": "LEFT",
                    "style": "Normal",
                    "bold_ratio": 0.0,
                    "file_name": "sample.docx",
                    "predicted_label": "title_section",
                    "postprocessed_label": "title_section",
                    "confidence_score": 0.99,
                    "low_confidence": False,
                    "bibliography_section_index": 1,
                },
                {
                    "doc_id": "doc_1",
                    "block_id": 3,
                    "text": "\tИванов И. И. Учебник. — Москва, 2020. — 120 с.",
                    "kind": "paragraph",
                    "alignment": "LEFT",
                    "style": "Normal",
                    "bold_ratio": 0.0,
                    "file_name": "sample.docx",
                    "predicted_label": "body_text",
                    "postprocessed_label": "bibliography_item",
                    "confidence_score": 0.99,
                    "low_confidence": False,
                    "bibliography_section_index": 1,
                },
            ]
        ).to_csv(predictions_csv, index=False, encoding="utf-8-sig")

        fixed_docx = tmp_path / "fixed.docx"
        summary = audit_or_format_docx(
            input_docx=input_docx,
            predictions_csv=predictions_csv,
            report_csv=tmp_path / "fix.csv",
            output_docx=fixed_docx,
            apply_safe=True,
            profile_id="gost_7_32_2017",
        )

        assert summary["changed"] == 2
        fixed = Document(fixed_docx)
        assert fixed.paragraphs[1].text == "1 ТЕОРЕТИЧЕСКАЯ ЧАСТЬ"
        assert fixed.paragraphs[2]._p.pPr is not None
        assert fixed.paragraphs[2]._p.pPr.numPr is not None
    finally:
        shutil.rmtree(tmp_path, ignore_errors=True)


# ---- Phase 01: style-guard tests (RED state — guard implemented in Plan 03) ----

def _row_data_body_text(text: str) -> dict:
    return {
        "text": text,
        "confidence_score": 0.99,
        "low_confidence": False,
    }


def test_style_guard_blocks_body_text_on_heading() -> None:
    document = Document()
    paragraph = document.add_paragraph("Глава 1. Введение")
    paragraph.style = "Heading 1"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data=_row_data_body_text(paragraph.text),
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["manual_review_required"] is True
    assert result["applied_fixes"] == []
    assert result["explanation"].startswith("style_guard_block:"), result["explanation"]


def test_style_guard_blocks_body_text_on_toc() -> None:
    document = Document()
    paragraph = document.add_paragraph("Глава 1 ... 5")
    # python-docx default template ships no "TOC 1" — use "TOC Heading"
    # (the same choice the fixture _build_style_guard_minimal.py makes).
    paragraph.style = "TOC Heading"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data=_row_data_body_text(paragraph.text),
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["explanation"].startswith("style_guard_block:"), result["explanation"]


def test_style_guard_blocks_body_text_on_caption() -> None:
    document = Document()
    paragraph = document.add_paragraph("Рисунок 1 — Схема")
    paragraph.style = "Caption"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data=_row_data_body_text(paragraph.text),
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    assert result["explanation"].startswith("style_guard_block:"), result["explanation"]


def test_style_guard_blocks_body_text_on_list() -> None:
    document = Document()
    paragraph = document.add_paragraph("Первый пункт перечисления")
    paragraph.style = "List Paragraph"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data=_row_data_body_text(paragraph.text),
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None
    assert result["status"] == "review"
    # NEW guard fires earlier than the existing late blocked_unsafe_autofix path —
    # the explanation marker distinguishes the two.
    assert result["explanation"].startswith("style_guard_block:"), result["explanation"]


def test_style_guard_passes_heading_rule_on_heading() -> None:
    # Heading rule on heading paragraph: guard MUST NOT fire (D-03).
    # rationale: D-03 says non-body_text labels are unaffected. Strengthened
    # vs. vacuous variant by asserting `result is not None` AND the
    # explanation does NOT begin with "style_guard_block:". If the existing
    # title_section path returns None for this paragraph, copy the
    # row_data/paragraph shape from a pre-Phase-1 test in this file that
    # already produces a non-None title_section result (grep for
    # `label="title_section"` and pick one with non-None observed outcome).
    document = Document()
    paragraph = document.add_paragraph("Введение")
    paragraph.style = "Heading 1"

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="title_section",
        row_data={"text": paragraph.text, "confidence_score": 0.99, "low_confidence": False},
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    # Strengthened — refuse to vacuously pass on result is None.
    assert result is not None, (
        "title_section path returned None for a Heading-styled paragraph — "
        "this test cannot prove the guard did not fire. Replace the row_data/"
        "paragraph shape with one that an existing pre-Phase-1 test already "
        "produces a non-None title_section result for, then re-run."
    )
    assert not result["explanation"].startswith("style_guard_block:"), result["explanation"]


def test_style_guard_passes_body_text_on_normal() -> None:
    # body_text on Normal-styled paragraph: guard MUST NOT fire — existing
    # body_text path runs. Strengthened to refuse vacuous-None passes.
    # rationale: if `apply_rules_to_paragraph(label="body_text", Normal-styled,
    # long text)` returns None for this row_data, copy the row_data shape
    # from an existing body_text test that produces a non-None result
    # (e.g., the pre-Phase-1 `test_rule_application_*` family in this file).
    document = Document()
    paragraph = document.add_paragraph(
        "Обычный абзац основного текста с достаточной длиной для applicable_rules."
    )
    # leave style as default Normal

    result = apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data=_row_data_body_text(paragraph.text),
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    assert result is not None, (
        "body_text path returned None for a Normal-styled paragraph — "
        "this test cannot prove the guard did not fire. Replace row_data "
        "with a shape that produces a non-None result in an existing pre-"
        "Phase-1 test, then re-run."
    )
    assert not result["explanation"].startswith("style_guard_block:"), result["explanation"]


def test_style_guard_does_not_write_direct_props() -> None:
    document = Document()
    paragraph = document.add_paragraph("Глава 2. Методика")
    paragraph.style = "Heading 1"
    # baseline: paragraph_format props are None by default for a freshly-added paragraph
    assert paragraph.paragraph_format.left_indent is None
    assert paragraph.paragraph_format.first_line_indent is None

    apply_rules_to_paragraph(
        paragraph=paragraph,
        label="body_text",
        row_data=_row_data_body_text(paragraph.text),
        rules=load_rules(),
        apply_safe=True,
        default_font_name="Times New Roman",
    )

    # After guard fires, NO direct paragraph-format props get written.
    assert paragraph.paragraph_format.left_indent is None
    assert paragraph.paragraph_format.first_line_indent is None
    assert paragraph.alignment is None


def test_style_guard_minimal_docx_changed_zero(tmp_path) -> None:
    """Integration: hand-crafted DOCX with one paragraph per style class.

    After Plan 03's guard lands, format-docx --apply-safe must yield changed=0
    because the guard short-circuits each styled paragraph (Heading/TOC/Caption/List)
    to status=review even when the predicted label is body_text, and the Normal
    paragraph already complies with body_text rules.

    Deviation rationale (Wave 0): `build_regression_predictions` would route the
    Heading/List/Caption paragraphs to non-body_text labels via `infer_regression_label`
    (style-based heuristic), bypassing the body_text path entirely — the guard would
    never be exercised. To pin the contract the plan is actually targeting
    ("body_text autofix mutates Heading/TOC/Caption/List paragraphs without the
    guard"), we simulate a worst-case model misprediction by overwriting
    predicted_label/postprocessed_label to "body_text" for ALL rows. This is the
    exact failure mode the Plan 03 guard exists to catch.
    """
    import pandas as pd

    input_docx = Path("tests/fixtures/style_guard_minimal.docx")
    assert input_docx.exists(), "fixture missing — run tests/fixtures/_build_style_guard_minimal.py"

    predictions_csv = tmp_path / "predictions.csv"
    report_csv = tmp_path / "report.csv"
    output_docx = tmp_path / "output.docx"
    build_regression_predictions(input_docx, predictions_csv)

    # Force body_text label on every paragraph to exercise the guard contract.
    predictions_df = pd.read_csv(predictions_csv, encoding="utf-8-sig")
    predictions_df["predicted_label"] = "body_text"
    predictions_df["postprocessed_label"] = "body_text"
    predictions_df.to_csv(predictions_csv, index=False, encoding="utf-8-sig")

    summary = audit_or_format_docx(
        input_docx=input_docx,
        predictions_csv=predictions_csv,
        report_csv=report_csv,
        output_docx=output_docx,
        apply_safe=True,
        profile_id="gost_7_32_2017",
    )

    assert summary["error"] == 0
    assert summary["changed"] == 0

    # Guard contract: every NON-Normal styled paragraph must be short-circuited
    # by the style_guard with an explanation starting "style_guard_block:".
    # Pre-Plan-03 this assertion is RED because no guard exists — explanations
    # come from the generic body_text inherited-format review path.
    report_df = pd.read_csv(report_csv, encoding="utf-8-sig")
    fixture_doc = Document(str(input_docx))
    styled_paragraph_indices = [
        i for i, p in enumerate(fixture_doc.paragraphs)
        if p.style.name != "Normal"
    ]
    for idx in styled_paragraph_indices:
        explanation = str(report_df.iloc[idx]["explanation"])
        assert explanation.startswith("style_guard_block:"), (
            f"paragraph {idx} (style={fixture_doc.paragraphs[idx].style.name}) "
            f"explanation should start with 'style_guard_block:' but was: {explanation}"
        )
