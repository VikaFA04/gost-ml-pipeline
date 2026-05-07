from __future__ import annotations
from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.generate.docx_styles import (
    configure_document_styles,
    setup_document_page,
)
from src.rules.docx_style_profile import (
    apply_style_profile_to_paragraph,
    get_default_font_name,
    get_label_style_profile,
    resolve_writer_style_name,
    resolve_writer_template_path,
)
from src.rules.profile_loader import load_profile

LABEL_COL_CANDIDATES = ["postprocessed_label", "predicted_label"]

def resolve_label_column(df: pd.DataFrame) -> str:
    for col in LABEL_COL_CANDIDATES:
        if col in df.columns:
            return col
    raise ValueError("В DataFrame нет ни 'postprocessed_label', ни 'predicted_label'")


def add_styled_paragraph(document: Document, text: str, label: str, profile: dict | None = None):
    paragraph = document.add_paragraph()
    style_name = resolve_writer_style_name(profile, label)
    try:
        paragraph.style = style_name
    except Exception:
        paragraph.style = "Normal"

    paragraph.add_run(text)

    style_profile = get_label_style_profile(profile, label)
    if style_profile:
        apply_style_profile_to_paragraph(
            paragraph=paragraph,
            style_profile=style_profile,
            default_font_name=get_default_font_name(profile),
        )

    return paragraph

def parse_table_text(table_text: str) -> list[list[str]]:
    """
    Восстанавливает таблицу из текстового представления:
    строки разделены '\\n', ячейки — ' | '.
    """
    rows = []
    for line in str(table_text).splitlines():
        line = line.strip()
        if not line:
            continue
        cells = [cell.strip() for cell in line.split(" | ")]
        rows.append(cells)
    return rows


def add_table_from_text(document: Document, table_text: str) -> None:
    rows = parse_table_text(table_text)
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j in range(max_cols):
            value = row[j] if j < len(row) else ""
            cell_par = table.cell(i, j).paragraphs[0]
            cell_par.add_run(value)
            cell_par.alignment = WD_ALIGN_PARAGRAPH.LEFT


def generate_docx_from_predictions(
    predictions_csv: str | Path,
    output_docx: str | Path,
    profile_path: str | Path | None = None,
    profile_id: str | None = None,
) -> None:
    predictions_csv = Path(predictions_csv)
    output_docx = Path(output_docx)
    profile = load_profile(profile_path=profile_path, profile_id=profile_id)

    df = pd.read_csv(predictions_csv)
    if "block_id" in df.columns:
        df = df.sort_values("block_id").reset_index(drop=True)

    label_col = resolve_label_column(df)

    template_path = resolve_writer_template_path(profile)
    if template_path is not None:
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон DOCX не найден: {template_path}")
        document = Document(str(template_path))
    else:
        document = Document()
    setup_document_page(document)
    configure_document_styles(document)

    for _, row in df.iterrows():
        text = str(row.get("text", "")).strip()
        kind = str(row.get("kind", "paragraph"))
        label = str(row.get(label_col, "body_text"))

        if not text:
            continue

        if kind == "table":
            add_table_from_text(document, text)
            continue

        add_styled_paragraph(document, text, label, profile=profile)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
