from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Callable

import pandas as pd
from docx import Document
from docx.table import Table

from src.evaluation.docx_style_diff import DocxStyleDiff, compare_docx_styles
from src.generate.inplace_formatter import audit_or_format_docx, extract_table_text
from src.io.block_extractor import extract_blocks_from_docx
from src.postprocess.postprocess_rules import apply_postprocess_rules


WORD_RE = re.compile(r"[A-Za-zА-Яа-яЁё0-9_]+")
BIBLIOGRAPHY_NUMBERED_SUBHEADING_RE = re.compile(
    r"^\d+\s*(теоретическая\s+часть|практическая\s+часть)$",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class NegativePairAudit:
    positive_path: Path
    negative_path: Path
    formatted_path: Path
    similarity: float
    before: DocxStyleDiff
    after: DocxStyleDiff
    formatter_summary: dict[str, object]

    @property
    def diff_delta(self) -> float:
        return self.after.diff_rate - self.before.diff_rate


def infer_regression_label(row: pd.Series) -> str:
    style = str(row.get("style") or "").lower()
    text = str(row.get("text") or "").strip().lower()

    if row.get("kind") == "table":
        return "table"
    if "heading 1" in style:
        return "title_section"
    if "heading 2" in style or "heading 3" in style:
        return "title_subsection"
    if BIBLIOGRAPHY_NUMBERED_SUBHEADING_RE.search(text):
        return "title_section"
    if "list" in style or pd.notna(row.get("list_type")):
        return "list_item"
    if text.startswith("таблица"):
        return "table_caption"
    if text.startswith("рисунок"):
        return "figure_caption"
    return "body_text"


def build_regression_predictions(input_docx: Path, predictions_csv: Path) -> None:
    df = extract_blocks_from_docx(input_docx)
    labels = [infer_regression_label(row) for _, row in df.iterrows()]
    df["predicted_label"] = labels
    df = apply_postprocess_rules(df, pred_col="predicted_label", out_col="postprocessed_label")
    df["confidence_score"] = 0.99
    df["low_confidence"] = False
    df.to_csv(predictions_csv, index=False, encoding="utf-8-sig")


def extract_docx_words(path: Path) -> set[str]:
    document = Document(path)
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            chunks.append(paragraph.text)
    for table in document.tables:
        text = extract_table_text(table)
        if text:
            chunks.append(text)
    return {word.lower() for word in WORD_RE.findall(" ".join(chunks))}


def text_jaccard(left_words: set[str], right_words: set[str]) -> float:
    if not left_words and not right_words:
        return 1.0
    union = left_words | right_words
    if not union:
        return 0.0
    return len(left_words & right_words) / len(union)


def best_positive_match(negative_path: Path, positive_paths: list[Path]) -> tuple[Path, float]:
    negative_words = extract_docx_words(negative_path)
    scored = [
        (text_jaccard(negative_words, extract_docx_words(positive_path)), positive_path)
        for positive_path in positive_paths
    ]
    similarity, positive_path = max(scored, key=lambda item: item[0])
    return positive_path, similarity


def audit_negative_pair(
    positive_path: Path,
    negative_path: Path,
    workspace_dir: Path,
    profile_id: str = "gost_7_32_2017",
) -> NegativePairAudit:
    workspace_dir.mkdir(parents=True, exist_ok=True)
    predictions_csv = workspace_dir / f"{negative_path.stem}_predictions.csv"
    report_csv = workspace_dir / f"{negative_path.stem}_report.csv"
    formatted_path = workspace_dir / f"{negative_path.stem}_safe_formatted.docx"

    build_regression_predictions(negative_path, predictions_csv)
    summary = audit_or_format_docx(
        input_docx=negative_path,
        predictions_csv=predictions_csv,
        report_csv=report_csv,
        output_docx=formatted_path,
        apply_safe=True,
        profile_id=profile_id,
    )

    before = compare_docx_styles(positive_path, negative_path)
    after = compare_docx_styles(positive_path, formatted_path)
    similarity = text_jaccard(extract_docx_words(positive_path), extract_docx_words(negative_path))

    return NegativePairAudit(
        positive_path=positive_path,
        negative_path=negative_path,
        formatted_path=formatted_path,
        similarity=similarity,
        before=before,
        after=after,
        formatter_summary=summary,
    )


def audit_negative_directory(
    positive_dir: Path,
    negative_dir: Path,
    workspace_dir: Path,
    profile_id: str = "gost_7_32_2017",
    limit: int | None = None,
    progress_callback: Callable[[int, int, Path], None] | None = None,
) -> list[NegativePairAudit]:
    positive_paths = sorted(path for path in positive_dir.glob("*.docx") if not path.name.startswith("~$"))
    audits: list[NegativePairAudit] = []
    negative_paths = sorted(path for path in negative_dir.glob("*.docx") if not path.name.startswith("~$"))
    if limit is not None:
        negative_paths = negative_paths[: max(limit, 0)]
    total = len(negative_paths)
    for index, negative_path in enumerate(negative_paths, start=1):
        if progress_callback is not None:
            progress_callback(index, total, negative_path)
        positive_path, _ = best_positive_match(negative_path, positive_paths)
        pair_workspace = workspace_dir / negative_path.stem
        audits.append(
            audit_negative_pair(
                positive_path=positive_path,
                negative_path=negative_path,
                workspace_dir=pair_workspace,
                profile_id=profile_id,
            )
        )
    return audits


def audits_to_frame(audits: list[NegativePairAudit]) -> pd.DataFrame:
    rows = []
    for audit in audits:
        before_field_mismatches = sum(audit.before.field_mismatches.values())
        after_field_mismatches = sum(audit.after.field_mismatches.values())
        rows.append(
            {
                "positive": audit.positive_path.name,
                "negative": audit.negative_path.name,
                "formatted": str(audit.formatted_path),
                "text_similarity": round(audit.similarity, 6),
                "before_diff_rate": round(audit.before.diff_rate, 6),
                "after_diff_rate": round(audit.after.diff_rate, 6),
                "diff_delta": round(audit.diff_delta, 6),
                "before_changed": audit.before.changed_paragraphs,
                "after_changed": audit.after.changed_paragraphs,
                "before_field_mismatches": before_field_mismatches,
                "after_field_mismatches": after_field_mismatches,
                "field_mismatch_delta": after_field_mismatches - before_field_mismatches,
                "formatter_changed": audit.formatter_summary.get("changed", 0),
                "formatter_review": audit.formatter_summary.get("review", 0),
                "formatter_blocked_unsafe": audit.formatter_summary.get("blocked_unsafe_autofix", 0),
                "formatter_error": audit.formatter_summary.get("error", 0),
            }
        )
    return pd.DataFrame(rows)
