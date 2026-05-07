from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph


ALIGNMENT_NAMES = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "DISTRIBUTE",
}

STYLE_FIELDS = [
    "style",
    "alignment",
    "first_line_indent_cm",
    "left_indent_cm",
    "line_spacing",
    "space_before_pt",
    "space_after_pt",
    "font_size_pt",
    "bold",
]


@dataclass(frozen=True)
class DocxStyleDiff:
    positive_path: Path
    candidate_path: Path
    compared_paragraphs: int
    changed_paragraphs: int
    extra_positive_paragraphs: int
    extra_candidate_paragraphs: int
    field_mismatches: dict[str, int]

    @property
    def diff_rate(self) -> float:
        if self.compared_paragraphs == 0:
            return 0.0
        return self.changed_paragraphs / self.compared_paragraphs


def _round_measure(value: Any, unit: str) -> float | None:
    if value is None:
        return None
    raw = getattr(value, unit, None)
    if raw is None:
        return None
    return round(float(raw), 3)


def _alignment_name(value: Any) -> str | None:
    if value is None:
        return None
    return ALIGNMENT_NAMES.get(value, str(value))


def _first_text_run_signature(paragraph: Paragraph) -> dict[str, Any]:
    for run in paragraph.runs:
        if run.text and run.text.strip():
            font_size = run.font.size
            return {
                "font_size_pt": round(float(font_size.pt), 3) if font_size is not None else None,
                "bold": bool(run.bold) if run.bold is not None else False,
            }
    return {"font_size_pt": None, "bold": False}


def paragraph_style_signature(paragraph: Paragraph) -> dict[str, Any]:
    fmt = paragraph.paragraph_format
    line_spacing = fmt.line_spacing
    run_style = _first_text_run_signature(paragraph)

    return {
        "style": paragraph.style.name if paragraph.style is not None else None,
        "alignment": _alignment_name(paragraph.alignment),
        "first_line_indent_cm": _round_measure(fmt.first_line_indent, "cm"),
        "left_indent_cm": _round_measure(fmt.left_indent, "cm"),
        "line_spacing": round(float(line_spacing), 3) if isinstance(line_spacing, (int, float)) else None,
        "space_before_pt": _round_measure(fmt.space_before, "pt"),
        "space_after_pt": _round_measure(fmt.space_after, "pt"),
        "font_size_pt": run_style["font_size_pt"],
        "bold": run_style["bold"],
    }


def compare_docx_styles(positive_path: Path, candidate_path: Path) -> DocxStyleDiff:
    positive_doc = Document(positive_path)
    candidate_doc = Document(candidate_path)
    positive_signatures = [
        paragraph_style_signature(paragraph)
        for paragraph in positive_doc.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]
    candidate_signatures = [
        paragraph_style_signature(paragraph)
        for paragraph in candidate_doc.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]

    compared = min(len(positive_signatures), len(candidate_signatures))
    field_mismatches: Counter[str] = Counter()
    changed_paragraphs = 0

    for positive_signature, candidate_signature in zip(positive_signatures, candidate_signatures):
        changed = False
        for field in STYLE_FIELDS:
            if positive_signature.get(field) != candidate_signature.get(field):
                field_mismatches[field] += 1
                changed = True
        if changed:
            changed_paragraphs += 1

    return DocxStyleDiff(
        positive_path=Path(positive_path),
        candidate_path=Path(candidate_path),
        compared_paragraphs=compared,
        changed_paragraphs=changed_paragraphs,
        extra_positive_paragraphs=max(0, len(positive_signatures) - len(candidate_signatures)),
        extra_candidate_paragraphs=max(0, len(candidate_signatures) - len(positive_signatures)),
        field_mismatches=dict(field_mismatches),
    )
